"""PDF and Word document text extraction for the read tool.

Both formats store text inside a binary container the read tool cannot
show as lines on its own, so this pulls the text out first and hands
back plain text that the normal line-numbering and pagination in
read_tool can treat like any other file.
"""

from pathlib import Path
from typing import Union

DOCUMENT_EXTENSIONS = {".pdf", ".docx", ".doc"}


def is_document_path(path: Path) -> bool:
    """True when `path` is a format this module knows how to extract."""

    return path.suffix.lower() in DOCUMENT_EXTENSIONS


def extract_document_text(
    path: Path,
) -> tuple[Union[str, None], str]:  # noqa: UP007, RUF100
    """Extract plain text from a PDF or Word document.

    Returns `(text, "")` on success, or `(None, reason)` explaining why
    it could not be read.
    """

    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return _extract_pdf(path)
    if suffix == ".docx":
        return _extract_docx(path)
    if suffix == ".doc":
        return None, (
            f"{path} is a legacy .doc file, which this tool cannot parse. "
            "Save it as .docx and read that instead."
        )
    return None, f"{path} is not a document type this tool can read."


def _extract_pdf(
    path: Path,
) -> tuple[Union[str, None], str]:  # noqa: UP007, RUF100
    try:
        from pypdf import PdfReader
    except ImportError:
        return None, (
            "pypdf is not installed, so PDF files cannot be read. "
            "Install it with: pip install pypdf"
        )

    try:
        reader = PdfReader(str(path))
    except Exception as exc:  # noqa: BLE001
        return None, f"Could not open {path} as a PDF: {exc}"

    if reader.is_encrypted:
        return None, f"{path} is password-protected and cannot be read."

    pages = []
    any_text = False
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            any_text = True
            pages.append(f"--- page {index} ---\n{text}")
        else:
            pages.append(f"--- page {index} (no extractable text) ---")

    if not any_text:
        return None, (
            f"No extractable text in {path}. It may be a scanned or "
            "image-only PDF."
        )

    return "\n\n".join(pages), ""


def _extract_docx(
    path: Path,
) -> tuple[Union[str, None], str]:  # noqa: UP007, RUF100
    try:
        import docx
    except ImportError:
        return None, (
            "python-docx is not installed, so Word documents cannot be "
            "read. Install it with: pip install python-docx"
        )

    try:
        document = docx.Document(str(path))
    except Exception as exc:  # noqa: BLE001
        return None, f"Could not open {path} as a Word document: {exc}"

    parts = [paragraph.text for paragraph in document.paragraphs]
    for table in document.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))

    combined = "\n".join(parts)
    if not combined.strip():
        return None, f"No extractable text in {path}."

    return combined, ""
