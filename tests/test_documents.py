# pylint: disable=C0114,C0115,C0116

import docx

from flash.documents import extract_document_text, is_document_path


def test_is_document_path(tmp_path):
    assert is_document_path(tmp_path / "report.pdf")  # nosec B101
    assert is_document_path(tmp_path / "report.DOCX")  # nosec B101
    assert not is_document_path(tmp_path / "report.txt")  # nosec B101


def test_extracts_docx_paragraphs_and_tables(tmp_path):
    target = tmp_path / "report.docx"
    document = docx.Document()
    document.add_paragraph("First paragraph.")
    document.add_paragraph("Second paragraph.")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "left"
    table.rows[0].cells[1].text = "right"
    document.save(target)

    text, reason = extract_document_text(target)

    assert reason == ""  # nosec B101
    assert text is not None  # nosec B101
    assert "First paragraph." in text  # nosec B101
    assert "left\tright" in text  # nosec B101


def test_empty_docx_reports_no_text(tmp_path):
    target = tmp_path / "empty.docx"
    docx.Document().save(target)

    text, reason = extract_document_text(target)

    assert text is None  # nosec B101
    assert "No extractable text" in reason  # nosec B101


def test_legacy_doc_is_rejected(tmp_path):
    target = tmp_path / "old.doc"
    target.write_bytes(b"not a real .doc file")

    text, reason = extract_document_text(target)

    assert text is None  # nosec B101
    assert ".docx" in reason  # nosec B101


def test_extracts_pdf_text_per_page(monkeypatch, tmp_path):
    target = tmp_path / "report.pdf"
    target.write_bytes(b"%PDF-1.4 fake")

    class FakePage:
        def __init__(self, text):
            self._text = text

        def extract_text(self):
            return self._text

    class FakeReader:
        is_encrypted = False

        def __init__(self, _path):
            self.pages = [FakePage("Page one text."), FakePage("")]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)

    text, reason = extract_document_text(target)

    assert reason == ""  # nosec B101
    assert text is not None  # nosec B101
    assert "--- page 1 ---\nPage one text." in text  # nosec B101
    assert "--- page 2 (no extractable text) ---" in text  # nosec B101


def test_encrypted_pdf_is_rejected(monkeypatch, tmp_path):
    target = tmp_path / "locked.pdf"
    target.write_bytes(b"%PDF-1.4 fake")

    class FakeReader:
        is_encrypted = True

        def __init__(self, _path):
            self.pages = []

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)

    text, reason = extract_document_text(target)

    assert text is None  # nosec B101
    assert "password-protected" in reason  # nosec B101


def test_scanned_pdf_with_no_text_is_rejected(monkeypatch, tmp_path):
    target = tmp_path / "scanned.pdf"
    target.write_bytes(b"%PDF-1.4 fake")

    class FakePage:
        def extract_text(self):
            return ""

    class FakeReader:
        is_encrypted = False

        def __init__(self, _path):
            self.pages = [FakePage()]

    monkeypatch.setattr("pypdf.PdfReader", FakeReader)

    text, reason = extract_document_text(target)

    assert text is None  # nosec B101
    assert "scanned or" in reason  # nosec B101
